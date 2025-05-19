# This script converts a markdown file to a Word document using the python-docx library.

# Import necessary libraries
import os
import re
import sys
from docx import Document


def read_markdown_file(file_path):
    """
    Read the markdown file and return its content.
    """
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


# Convert the markdown text to a Word document
def convert_markdown_to_word(markdown_text, output_file):
    """
    Convert markdown text to a Word document.
    """
    # Create a new Word document
    doc = Document()

    # Split the markdown text into lines
    lines = markdown_text.split("\n")

    # Process each line and this first line should be the title of the document and should be in the center
    title = lines[0]
    doc.add_heading(title, level=0).alignment = 1  # Center alignment
    doc.add_paragraph()  # Add a blank line after the title
    # Process each line in the markdown text
    for line in lines[1:]:
        # Check for headers
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("##### "):
            doc.add_heading(line[6:], level=5)
        elif line.startswith("###### "):
            doc.add_heading(line[7:], level=6)
        # Check for bold text
        elif "**" in line:
            bold_text = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            doc.add_paragraph(bold_text, style="ListBullet")
        # Check for italic text
        elif "*" in line:
            italic_text = re.sub(r"\*(.*?)\*", r"\1", line)
            doc.add_paragraph(italic_text, style="ListBullet")
        # Check for links
        elif "[" in line and "]" in line and "(" in line and ")" in line:
            link_text = re.search(r"\[(.*?)\]", line).group(1)
            link_url = re.search(r"\((.*?)\)", line).group(1)
            doc.add_paragraph(f"{link_text}: {link_url}", style="ListBullet")
        # Check for lists
        elif line.startswith("- ") or line.startswith("* "):
            list_item = line[2:]
            doc.add_paragraph(list_item, style="ListBullet")
        # Check for lists with numbers
        elif re.match(r"^\d+\.", line):
            list_item = re.sub(r"^\d+\.\s*", "", line)
            doc.add_paragraph(list_item, style="ListNumber")
        # Add regular text
        else:
            doc.add_paragraph(line)

    # Add page number to footer (centered)
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = 1  # Center alignment
    run = paragraph.add_run()
    fldChar1 = run._element
    fldChar1.text = ""
    # Insert PAGE field
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Save the Word document
    doc.save(output_file)


# Read the markdown file "Project Report: Market Analysis Q1 2025.md"
markdown_text = read_markdown_file("Project Report: Market Analysis Q1 2025.md")

# Convert the markdown text to a Word document
convert_markdown_to_word(markdown_text, "Project Report: Market Analysis Q1 2025.docx")
